#!/usr/bin/env python3
"""Replace ASCII box-drawing figures in JFAR docx with Graphviz PNG exports."""

from __future__ import annotations

import re
import shutil
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.shared import Inches
from docx.text.paragraph import Paragraph

ROOT = Path(__file__).resolve().parents[1]
SRC = ROOT / "JFAR Revision-August 2026-Ron_04Jul26.docx"
OUT = ROOT / "JFAR Revision-August 2026-Ron_04Jul26-graphics.docx"
FIGURES = Path(__file__).resolve().parent / "figures"

BOX_CHARS = set("┌┐└┘├┤│─▼▲◄►┬┴┼")


def is_ascii_figure_line(text: str) -> bool:
    t = text.strip()
    if not t:
        return False
    if "Training Flow A" in t or "Training Flow B" in t or "Training Flow C" in t:
        return True
    if "Fraud Analysis A" in t or "Fraud Analysis B" in t or "Fraud Analysis C" in t:
        return True
    return any(ch in t for ch in BOX_CHARS)


def delete_paragraph(paragraph: Paragraph) -> None:
    element = paragraph._element
    element.getparent().remove(element)
    paragraph._p = paragraph._element = None  # type: ignore[attr-defined]


def insert_paragraph_after(paragraph: Paragraph, style: str | None = None) -> Paragraph:
    new_p = OxmlElement("w:p")
    paragraph._p.addnext(new_p)
    new_para = Paragraph(new_p, paragraph._parent)
    if style:
        new_para.style = style
    return new_para


def add_centered_picture(
    after: Paragraph,
    image_path: Path,
    width_inches: float,
) -> Paragraph:
    para = insert_paragraph_after(after, style="Body A")
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = para.add_run()
    run.add_picture(str(image_path), width=Inches(width_inches))
    return para


def add_text_after(after: Paragraph, text: str, style: str = "Body A") -> Paragraph:
    para = insert_paragraph_after(after, style=style)
    para.add_run(text)
    return para


def find_paragraph(doc: Document, needle: str) -> Paragraph:
    for para in doc.paragraphs:
        if needle in para.text:
            return para
    raise ValueError(f"Paragraph not found containing: {needle!r}")


def remove_ascii_blocks(doc: Document) -> int:
    removed = 0
    for para in list(doc.paragraphs):
        if is_ascii_figure_line(para.text):
            delete_paragraph(para)
            removed += 1
    return removed


def update_caption(doc: Document, old_prefix: str, new_text: str) -> None:
    for para in doc.paragraphs:
        if para.text.strip().startswith(old_prefix):
            para.clear()
            para.add_run(new_text)
            return


def main() -> None:
    if not SRC.exists():
        raise SystemExit(f"Source not found: {SRC}")

    shutil.copy2(SRC, OUT)
    doc = Document(str(OUT))

    removed = remove_ascii_blocks(doc)
    print(f"Removed {removed} ASCII figure paragraphs")

    # Figure 4 — MS-readable 2x2 macro layout (large type)
    anchor4 = find_paragraph(doc, "Figure 4 depicts the high-level agentive architecture")
    last = add_centered_picture(
        anchor4,
        FIGURES / "Figure4_agent_network_MS.png",
        width_inches=6.5,
    )
    update_caption(doc, "Figure 4.", "Figure 4. NGAME Agent Interaction Network")

    # Figure 6 — portrait training flow (replaces A/B/C ASCII segments)
    anchor6 = find_paragraph(
        doc,
        "The following multi-page flow diagram depicts the Training Pipeline",
    )
    last = add_centered_picture(
        anchor6,
        FIGURES / "Figure6_training_flow.png",
        width_inches=5.5,
    )
    last = add_text_after(last, "Figure 6. NGAME Training Flow", style="Body Text")

    # Figure 7 — two panels (caption, then panel descriptions, then graphics)
    anchor7 = find_paragraph(
        doc,
        "If the AP-level coefficient exceeds the value of its counterpart in the reference set",
    )
    last = add_text_after(anchor7, "Figure 7. NGAME Daily Fraud Analysis", style="Body Text")
    last = add_text_after(last, "Panel A: Detection and triage (Layers 1–2).", style="Body Text")
    last = add_centered_picture(last, FIGURES / "Figure7_PanelA.png", width_inches=5.5)
    last = add_text_after(last, "Panel B: Interpretation and alerting (Layers 3–4).", style="Body Text")
    add_centered_picture(last, FIGURES / "Figure7_PanelB.png", width_inches=5.5)

    doc.save(str(OUT))
    print(f"Saved: {OUT}")


if __name__ == "__main__":
    main()
