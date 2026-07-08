#!/usr/bin/env python3
"""Swap Figure 4 in the graphics docx for the MS-readable variant."""

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches

ROOT = Path(__file__).resolve().parents[1]
DOCX = ROOT / "JFAR Revision-August 2026-Ron_04Jul26-graphics.docx"
FIGURE = Path(__file__).resolve().parent / "figures" / "Figure4_agent_network_MS.png"
ANCHOR = "Figure 4 depicts the high-level agentive architecture"
# Full text width; figure aspect ~1.88:1 → height ≈ 3.5 in at this width
WIDTH_INCHES = 6.5


def main() -> None:
    if not DOCX.exists():
        raise SystemExit(f"Not found: {DOCX}")
    if not FIGURE.exists():
        raise SystemExit(f"Not found: {FIGURE}")

    doc = Document(str(DOCX))
    anchor_idx = None
    for i, para in enumerate(doc.paragraphs):
        if ANCHOR in para.text:
            anchor_idx = i
            break
    if anchor_idx is None:
        raise SystemExit(f"Anchor paragraph not found: {ANCHOR!r}")

    image_idx = None
    for j in range(anchor_idx + 1, min(anchor_idx + 6, len(doc.paragraphs))):
        if doc.paragraphs[j]._element.xpath(".//a:blip"):
            image_idx = j
            break
    if image_idx is None:
        raise SystemExit("Figure 4 image paragraph not found after anchor")

    para = doc.paragraphs[image_idx]
    para.clear()
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    para.add_run().add_picture(str(FIGURE), width=Inches(WIDTH_INCHES))

    doc.save(str(DOCX))
    print(f"Updated Figure 4 in {DOCX}")


if __name__ == "__main__":
    main()
